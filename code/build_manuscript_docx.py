"""Build an illustrated Word version of the integration manuscript from
docs/manuscript_draft.md (markdown with inline ![](figure) embeds)."""
import os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__)); ROOT = os.path.dirname(HERE)
DOCS = os.path.join(ROOT, 'docs')
md = open(os.path.join(DOCS, 'manuscript_draft.md'), encoding='utf-8').read().splitlines()

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
for m in ('top', 'bottom', 'left', 'right'):
    setattr(sec, f'{m}_margin', Inches(1))
doc.styles['Normal'].font.name = 'Arial'; doc.styles['Normal'].font.size = Pt(11)

INLINE = re.compile(r'(\*\*.+?\*\*|\*.+?\*)')
IMG = re.compile(r'^!\[.*?\]\((.+?)\)\s*$')

def add_runs(par, text):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            par.add_run(tok[1:-1]).italic = True
        else:
            par.add_run(tok)

for line in md:
    s = line.rstrip()
    if not s.strip():
        continue
    mi = IMG.match(s)
    if mi:
        rel = mi.group(1)
        path = os.path.normpath(os.path.join(DOCS, rel))
        if os.path.exists(path):
            w, h = Image.open(path).size
            width = Inches(6.5 if h / w <= 1.1 else min(6.5, 8.3 * w / h))
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=width)
        continue
    if s.startswith('# '):
        add_runs(doc.add_heading(level=0), s[2:])
    elif s.startswith('### '):
        add_runs(doc.add_heading(level=2), s[4:])
    elif s.startswith('## '):
        add_runs(doc.add_heading(level=1), s[3:])
    elif s.startswith('- '):
        add_runs(doc.add_paragraph(style='List Bullet'), s[2:])
    else:
        p = doc.add_paragraph()
        if s.startswith('*') and s.endswith('*') and not s.startswith('**'):
            r = p.add_run(s.strip('*')); r.italic = True; r.font.size = Pt(9.5)  # figure caption
        else:
            add_runs(p, s)

# fix python-docx zoom-attr schema nit
from docx.oxml.ns import qn
z = doc.settings.element.find(qn('w:zoom'))
if z is not None and z.get(qn('w:percent')) is None:
    z.set(qn('w:percent'), '100')
out = os.path.join(DOCS, 'manuscript_with_figures.docx')
doc.save(out)
print('saved', os.path.getsize(out), 'bytes ->', out)
print('images embedded:', sum(1 for _ in doc.inline_shapes))
